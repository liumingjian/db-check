"""Render a ReportView into a template-styled DOCX."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

from reporter.model.errors import EXIT_OUTPUT_ERROR, EXIT_RENDER_ERROR, EXIT_TEMPLATE_ERROR, ReporterFailure
from reporter.renderer.inline_markup import append_inline_runs
from reporter.renderer.table_widths import apply_table_geometry, column_widths

HEADER_FILL = "DEEAF6"
HEADER_COLOR = RGBColor(0x18, 0x3D, 0x5E)
BODY_COLOR = RGBColor(0x33, 0x33, 0x33)
DOC_CONTROL_TITLE_SIZE = 18
DOC_CONTROL_SUBTITLE_SIZE = 16
CAPTION_SIZE = 10.5
TABLE_FONT_SIZE = 9
DOC_CONTROL_MARKER = "文档控制"
EMOJI_FONT = "Apple Color Emoji"
CHAPTER_PREFIX_PATTERN = re.compile(r"^第[一二三四五六七八九十百零0-9]+章\s*")
SECTION_PREFIX_PATTERN = re.compile(r"^\d+(?:\.\d+)*\s*")
RISK_ICON_COLUMNS = {"风险标识"}


def render_template_docx(template_path: Path, report_view_path: Path, output_path: Path) -> None:
    document = _load_template(template_path)
    report_view = _load_report_view(report_view_path)
    _clear_body_after_cover(document)
    chapter_state = {"top_level_count": 0}
    for section in report_view["sections"]:
        _render_section(document, section, level=0, doc_control=False, chapter_state=chapter_state)
    _save_document(document, output_path)


def _load_template(template_path: Path) -> Document:
    try:
        return Document(str(template_path))
    except Exception as exc:  # noqa: BLE001
        raise ReporterFailure(EXIT_TEMPLATE_ERROR, f"模板文件无效: {exc}") from exc


def _load_report_view(report_view_path: Path) -> dict[str, Any]:
    try:
        payload = json.loads(report_view_path.read_text(encoding="utf-8"))
    except Exception as exc:  # noqa: BLE001
        raise ReporterFailure(EXIT_RENDER_ERROR, f"report-view 读取失败: {exc}") from exc
    if not isinstance(payload, dict) or "sections" not in payload:
        raise ReporterFailure(EXIT_RENDER_ERROR, "report-view 结构无效")
    return payload


def _clear_body_after_cover(document: Document) -> None:
    body = document._body._element
    start_index = _find_content_start(body)
    for child in list(body)[start_index:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _find_content_start(body: Any) -> int:
    for index, child in enumerate(list(body)):
        if DOC_CONTROL_MARKER in _element_text(child):
            return index
    return 0


def _element_text(element: Any) -> str:
    return "".join(text.strip() for text in element.itertext())


def _render_section(
    document: Document,
    section: dict[str, Any],
    level: int,
    doc_control: bool,
    chapter_state: dict[str, int],
) -> None:
    title = str(section.get("title", ""))
    current_doc_control = doc_control or title == "文档控制"
    _maybe_add_page_break(document, level, current_doc_control, chapter_state)
    _add_heading(document, title, level, current_doc_control)
    note = str(section.get("note", ""))
    if note:
        _add_quote(document, f"状态: {_status_label(str(section.get('status', 'collected')))}")
        _add_quote(document, f"说明: {note}")
    for paragraph in section.get("paragraphs", []):
        _add_body_paragraph(document, str(paragraph))
    for table in section.get("tables", []):
        _render_table(document, table, title, note)
    for child in section.get("children", []):
        _render_section(document, child, level + 1, current_doc_control, chapter_state)


def _maybe_add_page_break(document: Document, level: int, doc_control: bool, chapter_state: dict[str, int]) -> None:
    if level != 0:
        return
    if doc_control:
        chapter_state["doc_control_rendered"] = 1
        return
    if chapter_state["top_level_count"] > 0 or chapter_state.get("doc_control_rendered", 0) > 0:
        document.add_page_break()
    chapter_state["doc_control_rendered"] = 0
    chapter_state["top_level_count"] += 1


def _add_heading(document: Document, title: str, level: int, doc_control: bool) -> None:
    if doc_control and level == 0:
        _add_custom_heading(document, title, DOC_CONTROL_TITLE_SIZE)
        return
    if doc_control and level == 1:
        _add_custom_heading(document, title, DOC_CONTROL_SUBTITLE_SIZE)
        return
    style_name = "Heading 1" if level == 0 else "Heading 2" if level == 1 else "Heading 3"
    document.add_paragraph(_heading_display_text(title), style=style_name)


def _heading_display_text(title: str) -> str:
    title = CHAPTER_PREFIX_PATTERN.sub("", title).strip()
    return SECTION_PREFIX_PATTERN.sub("", title).strip()


def _add_custom_heading(document: Document, title: str, size: int) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    _set_run_font(run, "黑体", size, bold=True)


def _add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    append_inline_runs(paragraph, text, "宋体", CAPTION_SIZE, BODY_COLOR, _set_run_font)


def _add_quote(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="markdown引用")
    paragraph.add_run(text)


def _render_table(document: Document, table_data: dict[str, Any], section_title: str, section_note: str) -> None:
    title = str(table_data.get("title", ""))
    if _should_render_caption(title, section_title):
        _add_table_caption(document, title)
    columns = tuple(str(item) for item in table_data.get("columns", []))
    rows = tuple(tuple("" if value is None else str(value) for value in row) for row in table_data.get("rows", []))
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.autofit = False
    _set_table_borders(table)
    widths = column_widths(document, columns, table_data.get("column_width_weights"))
    apply_table_geometry(table, widths)
    _fill_header(table.rows[0].cells, columns, widths)
    for row in rows:
        _fill_row(table.add_row().cells, columns, row, widths)
    if table_data.get("field_notes"):
        _add_field_notes(document, table_data["field_notes"])
    table_note = str(table_data.get("note", ""))
    if table_note and table_note != section_note:
        _add_quote(document, f"说明: {table_note}")
    document.add_paragraph()


def _add_table_caption(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title)
    _set_run_font(run, "黑体", CAPTION_SIZE, bold=True)


def _fill_header(cells: Any, columns: tuple[str, ...], widths: list[int]) -> None:
    for idx, cell in enumerate(cells):
        cell.width = Twips(widths[idx])
        _shade_cell(cell, HEADER_FILL)
        _set_cell_alignment(cell, columns[idx] in RISK_ICON_COLUMNS)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(columns[idx])
        _set_run_font(run, "Microsoft YaHei", TABLE_FONT_SIZE, bold=True, color=HEADER_COLOR)


def _fill_row(cells: Any, columns: tuple[str, ...], row: tuple[str, ...], widths: list[int]) -> None:
    for idx, cell in enumerate(cells):
        cell.width = Twips(widths[idx])
        use_emoji = columns[idx] in RISK_ICON_COLUMNS
        _set_cell_alignment(cell, use_emoji)
        paragraph = cell.paragraphs[0]
        font_name = EMOJI_FONT if use_emoji else "宋体"
        append_inline_runs(paragraph, row[idx], font_name, TABLE_FONT_SIZE, BODY_COLOR, _set_run_font, use_emoji)


def _add_field_notes(document: Document, field_notes: list[list[str]] | list[tuple[str, str]]) -> None:
    label = document.add_paragraph()
    run = label.add_run("字段说明：")
    _set_run_font(run, "黑体", CAPTION_SIZE, bold=True)
    for item in field_notes:
        if len(item) != 2:
            continue
        _add_quote(document, f"{item[0]}: {item[1]}")


def _should_render_caption(title: str, section_title: str) -> bool:
    return bool(title and title != _heading_display_text(section_title))


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "808080")


def _set_cell_alignment(cell: Any, centered: bool) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT


def _set_run_font(run: Any, font_name: str, size_pt: float, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color



def _status_label(status: str) -> str:
    labels = {
        "collected": "已采集",
        "derived": "已推导",
        "external": "外部补充",
        "missing": "未采集",
        "na": "不适用",
    }
    return labels.get(status, status)


def _save_document(document: Document, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    temp_path = output_path.with_suffix(output_path.suffix + ".tmp")
    try:
        document.save(str(temp_path))
        temp_path.replace(output_path)
    except ReporterFailure:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ReporterFailure(EXIT_OUTPUT_ERROR, f"DOCX 写入失败: {exc}") from exc
