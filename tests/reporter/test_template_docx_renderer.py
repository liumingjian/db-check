from __future__ import annotations

import json
import shutil
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from reporter.cli.render_template_docx import run
from reporter.content.mysql_report_builder import build_mysql_report_view

ROOT = Path(__file__).resolve().parents[2]


class TemplateDocxRendererTests(unittest.TestCase):
    def test_template_docx_cli_generates_styled_report(self) -> None:
        result = json.loads((ROOT / "contracts" / "result.sample.json").read_text(encoding="utf-8"))
        summary = json.loads((ROOT / "contracts" / "summary.sample.json").read_text(encoding="utf-8"))
        meta = json.loads((ROOT / "reporter" / "templates" / "report-meta.sample.json").read_text(encoding="utf-8"))
        report_view = build_mysql_report_view(result, summary, meta)
        temp_dir = Path(tempfile.mkdtemp())
        try:
            report_md = temp_dir / "report.md"
            report_md.write_text("# placeholder\n", encoding="utf-8")
            report_view_path = temp_dir / "report-view.json"
            report_view_path.write_text(json.dumps(report_view.to_dict(), ensure_ascii=False), encoding="utf-8")
            out_docx = temp_dir / "styled-report.docx"
            code = run(
                [
                    "--report-view",
                    str(report_view_path),
                    "--template",
                    str(ROOT / "reporter" / "templates" / "mysql-template.docx"),
                    "--out",
                    str(out_docx),
                ]
            )
            self.assertEqual(code, 0)
            self.assertTrue(out_docx.exists())
            document = Document(str(out_docx))
            texts = [paragraph.text for paragraph in document.paragraphs]
            self.assertIn("文档控制", texts)
            self.assertIn("巡检总结", texts)
            self.assertNotIn("第一章 巡检总结", texts)
            self.assertLess(texts.index("文档控制"), len(texts) - 1)
            self.assertEqual(texts[0], "")
            self.assertGreater(len(document.tables), 5)
            self.assertEqual(document.tables[0].style.name, "Table Grid")
            alarm_table = self._find_table(document, "风险等级", "风险标识", "定义", "建议响应时效")
            self.assertIsNotNone(alarm_table)
            alarm_widths = [cell.width for cell in alarm_table.rows[0].cells]
            self.assertGreater(alarm_widths[2], alarm_widths[0])
            self.assertGreater(alarm_widths[3], alarm_widths[1])
            health_table = self._find_table(document, "检查维度", "风险标识", "关键发现")
            self.assertIsNotNone(health_table)
            health_widths = [cell.width for cell in health_table.rows[0].cells]
            self.assertGreater(health_widths[2], health_widths[0])
            self.assertGreater(health_widths[2], health_widths[1])
            xml = self._document_xml(out_docx)
            self.assertIn("Apple Color Emoji", xml)
            self.assertIn('w:type="page"', xml)
            self.assertIn("w:tblBorders", xml)
            self.assertIn('<w:tblW w:type="dxa"', xml)
            self.assertIn('<w:gridCol w:w="996"/>', xml)
            self.assertIn('<w:gridCol w:w="664"/>', xml)
            self.assertIn('<w:gridCol w:w="4817"/>', xml)
            self.assertIn('<w:gridCol w:w="1829"/>', xml)
            self.assertIn('<w:gridCol w:w="1495"/>', xml)
            self.assertIn('<w:gridCol w:w="5815"/>', xml)
            self.assertIn("风险标识", "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells))
            system_table = self._find_table(document, "指标", "当前值", "说明")
            self.assertIsNotNone(system_table)
            value_by_label = {
                row.cells[0].text.strip(): row.cells[1].text.strip()
                for row in system_table.rows[1:]
                if len(row.cells) >= 2
            }
            self.assertEqual("52.10%", value_by_label.get("CPU 使用率"))
            self.assertEqual("3.40%", value_by_label.get("CPU iowait"))
            self.assertEqual("82.00%", value_by_label.get("内存使用率"))
            self.assertEqual("36.00%", value_by_label.get("磁盘使用率"))
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    def _document_xml(self, path: Path) -> str:
        with zipfile.ZipFile(path) as archive:
            return archive.read("word/document.xml").decode("utf-8")

    def _find_table(self, document: Document, *headers: str):
        expected = list(headers)
        for table in document.tables:
            current = [cell.text for cell in table.rows[0].cells]
            if current == expected:
                return table
        return None


if __name__ == "__main__":
    unittest.main()
